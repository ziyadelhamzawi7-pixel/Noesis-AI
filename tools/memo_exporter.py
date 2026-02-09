"""
Investment Memo DOCX Exporter

Converts memo data to a professionally formatted Word document.
"""

import re
import json
from io import BytesIO
from datetime import datetime
from typing import Dict, Any, Optional, List, Tuple

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from loguru import logger


# Section configuration
MEMO_SECTIONS = [
    ("proposed_investment_terms", "Proposed Investment Terms"),
    ("executive_summary", "Executive Summary"),
    ("market_analysis", "Market Analysis"),
    ("team_assessment", "Team Assessment"),
    ("product_technology", "Product & Technology"),
    ("financial_analysis", "Financial Analysis"),
    ("valuation_analysis", "Valuation Analysis"),
    ("risks_concerns", "Risks & Concerns"),
    ("outcome_scenario_analysis", "Outcome Scenario Analysis"),
    ("investment_recommendation", "Investment Recommendation"),
]


def _setup_styles(doc: Document) -> None:
    """Configure document styles for consistent formatting."""
    styles = doc.styles

    # Title style
    title_style = styles['Title']
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 0, 0)

    # Heading 1 style (section headers)
    h1_style = styles['Heading 1']
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 0, 0)
    h1_style.paragraph_format.space_before = Pt(24)
    h1_style.paragraph_format.space_after = Pt(12)

    # Heading 2 style (subsections)
    h2_style = styles['Heading 2']
    h2_style.font.size = Pt(13)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 0, 0)
    h2_style.paragraph_format.space_before = Pt(12)
    h2_style.paragraph_format.space_after = Pt(6)

    # Normal style (body text)
    normal_style = styles['Normal']
    normal_style.font.size = Pt(11)
    normal_style.font.name = 'Calibri'
    normal_style.paragraph_format.space_after = Pt(8)
    normal_style.paragraph_format.line_spacing = 1.15


def _add_cover_page(
    doc: Document,
    company_name: str,
    ticket_size: Optional[float],
    post_money_valuation: Optional[float],
    created_at: str
) -> None:
    """Add a cover page with company name and deal terms."""
    # Add some spacing at the top
    for _ in range(4):
        doc.add_paragraph()

    # Company name as title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(company_name)
    run.bold = True
    run.font.size = Pt(32)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Investment Memo")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    doc.add_paragraph()

    # Deal terms section
    if ticket_size or post_money_valuation:
        terms_header = doc.add_paragraph()
        terms_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = terms_header.add_run("Deal Terms")
        run.bold = True
        run.font.size = Pt(14)

        # Create a table for deal terms
        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if ticket_size:
            row = table.add_row()
            row.cells[0].text = "Ticket Size"
            row.cells[1].text = f"${ticket_size:,.0f}"

        if post_money_valuation:
            row = table.add_row()
            row.cells[0].text = "Post-Money Valuation"
            row.cells[1].text = f"${post_money_valuation:,.0f}"

        if ticket_size and post_money_valuation:
            ownership = (ticket_size / post_money_valuation) * 100
            row = table.add_row()
            row.cells[0].text = "Ownership"
            row.cells[1].text = f"{ownership:.2f}%"

        # Style the table
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    # Generation date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        date_obj = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
        date_str = date_obj.strftime("%B %d, %Y")
    except:
        date_str = created_at[:10] if created_at else datetime.now().strftime("%B %d, %Y")

    run = date_para.add_run(f"Generated: {date_str}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Page break after cover
    doc.add_page_break()


def _add_table_of_contents(doc: Document) -> None:
    """Add a table of contents page."""
    toc_title = doc.add_paragraph("Table of Contents", style='Heading 1')
    toc_title.paragraph_format.space_after = Pt(24)

    for i, (key, label) in enumerate(MEMO_SECTIONS, 1):
        para = doc.add_paragraph()
        run = para.add_run(f"{i}. {label}")
        run.font.size = Pt(12)
        para.paragraph_format.space_after = Pt(6)

    doc.add_page_break()


def _parse_markdown_table(lines: List[str], start: int) -> Tuple[List[str], List[List[str]], int]:
    """
    Parse a markdown table starting at the given line index.

    Returns:
        (headers, rows, end_index) where end_index is the first line after the table.
    """
    def _split_row(line: str) -> List[str]:
        # Strip leading/trailing pipes then split on |
        line = line.strip()
        if line.startswith('|'):
            line = line[1:]
        if line.endswith('|'):
            line = line[:-1]
        return [cell.strip() for cell in line.split('|')]

    headers = _split_row(lines[start])
    i = start + 1

    # Skip separator row (|---|---|)
    if i < len(lines) and re.match(r'^\s*\|[\s\-:|]+\|', lines[i]):
        i += 1

    rows = []
    while i < len(lines) and lines[i].strip().startswith('|'):
        rows.append(_split_row(lines[i]))
        i += 1

    return headers, rows, i


def _set_cell_shading(cell, color_hex: str) -> None:
    """Set background shading on a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_borders(cell) -> None:
    """Set thin borders on all sides of a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')  # 0.5pt
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'D2D2D2')
        borders.append(el)
    tc_pr.append(borders)


def _add_cell_content(cell, text: str, bold_all: bool = False, font_size: int = 10) -> None:
    """
    Populate a table cell, handling <br> tags and inline bold/italic markdown.
    """
    # Clear default empty paragraph
    cell.text = ''
    p = cell.paragraphs[0]

    # Split on <br> or <br/> for multi-line cells
    segments = re.split(r'<br\s*/?>', text)
    for seg_idx, segment in enumerate(segments):
        if seg_idx > 0:
            # Add a new paragraph for each <br>
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        segment = segment.strip()
        if not segment:
            continue
        _add_formatted_text(p, segment, default_size=Pt(font_size), default_bold=bold_all)


def _add_table_to_docx(doc: Document, headers: List[str], rows: List[List[str]]) -> None:
    """Create a styled DOCX table from parsed markdown table data."""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.autofit = True

    # Header row
    for col_idx, header_text in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        _add_cell_content(cell, header_text, bold_all=True)
        _set_cell_shading(cell, 'F0EEF8')  # matches --bg-tertiary
        _set_cell_borders(cell)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx in range(num_cols):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell_text = row_data[col_idx] if col_idx < len(row_data) else ''
            _add_cell_content(cell, cell_text)
            _set_cell_borders(cell)
            # Alternating row shading (even data rows)
            if row_idx % 2 == 1:
                _set_cell_shading(cell, 'F8F8FA')


def _markdown_to_docx(doc: Document, content: str) -> None:
    """
    Convert markdown content to Word paragraphs.
    Handles basic markdown: headers, bold, italic, bullets, numbered lists, tables.
    """
    if not content:
        return

    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Handle markdown tables
        if line.startswith('|'):
            headers, rows, i = _parse_markdown_table(lines, i)
            _add_table_to_docx(doc, headers, rows)
            continue

        # Handle headers (## Header)
        if line.startswith('### '):
            para = doc.add_paragraph(line[4:], style='Heading 2')
        elif line.startswith('## '):
            para = doc.add_paragraph(line[3:], style='Heading 2')
        elif line.startswith('# '):
            para = doc.add_paragraph(line[2:], style='Heading 1')

        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            para = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(para, line[2:])

        # Handle numbered lists
        elif re.match(r'^\d+\.\s', line):
            para = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\.\s', '', line)
            _add_formatted_text(para, text)

        # Regular paragraph
        else:
            para = doc.add_paragraph()
            _add_formatted_text(para, line)

        i += 1


def _add_formatted_text(para, text: str, default_size=None, default_bold: bool = False) -> None:
    """
    Add text with inline formatting (bold, italic) to a paragraph.
    Handles **bold**, *italic*, and ***bold italic***.

    Args:
        para: Paragraph to add runs to
        text: Markdown-formatted text
        default_size: Optional font size (Pt) to apply to all runs
        default_bold: If True, make all runs bold (used for table headers)
    """
    # Pattern to match bold/italic markdown
    pattern = r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|__.*?__|_.*?_)'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue

        if part.startswith('***') and part.endswith('***'):
            run = para.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('__') and part.endswith('__'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('_') and part.endswith('_'):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            run = para.add_run(part)
            if default_bold:
                run.bold = True

        if default_size:
            run.font.size = default_size


def _format_chart_value(value: float, fmt: str) -> str:
    """Format a value for chart display based on format type."""
    if fmt == 'currency':
        if abs(value) >= 1_000_000:
            return f'${value/1_000_000:.1f}M'
        if abs(value) >= 1_000:
            return f'${value/1_000:.0f}K'
        return f'${value:,.0f}'
    if fmt == 'percent':
        return f'{value:.0f}%'
    if abs(value) >= 1_000_000:
        return f'{value/1_000_000:.1f}M'
    if abs(value) >= 1_000:
        return f'{value/1_000:.0f}K'
    return f'{value:,.0f}'


def _generate_chart_images(chart_data: Dict[str, Any]) -> List[Tuple[str, BytesIO]]:
    """Generate matplotlib chart images from Claude-generated chart specs.

    Returns a list of (chart_id, BytesIO) tuples.
    """
    try:
        import matplotlib
        matplotlib.use('Agg')  # Non-interactive backend
        import matplotlib.pyplot as plt
        import matplotlib.ticker as mticker
    except ImportError:
        logger.warning("matplotlib not installed — skipping chart generation")
        return []

    charts = chart_data.get("charts", [])
    if not charts:
        return []

    images: List[Tuple[str, BytesIO]] = []

    for spec in charts:
        chart_id = spec.get("id", "chart")
        chart_type = spec.get("type", "bar")
        title = spec.get("title", "")
        x_key = spec.get("x_key", "")
        y_key = spec.get("y_key", "")
        y_format = spec.get("y_format", "number")
        color_key = spec.get("color_key", "")
        colors_config = spec.get("colors", "#6366f1")
        data = spec.get("data", [])

        if not data or not x_key or not y_key:
            continue

        try:
            x_values = [d.get(x_key, "") for d in data]
            y_values = [float(d.get(y_key, 0)) for d in data]

            # Determine bar colors
            if color_key and isinstance(colors_config, dict):
                bar_colors = [colors_config.get(str(d.get(color_key, "")), "#6366f1") for d in data]
            elif isinstance(colors_config, str):
                bar_colors = [colors_config] * len(data)
            else:
                bar_colors = ["#6366f1"] * len(data)

            is_horizontal = chart_type == "horizontal_bar"
            is_line = chart_type == "line"
            fig_height = max(2.5, 0.6 * len(data)) if is_horizontal else 3.5
            fig, ax = plt.subplots(figsize=(6, fig_height))

            if is_line:
                line_color = bar_colors[0] if bar_colors else "#6366f1"
                ax.plot(x_values, y_values, marker='o', color=line_color, linewidth=2, markersize=5)
                ax.set_ylabel(spec.get("y_label", ""), fontsize=10)
                ax.yaxis.set_major_formatter(mticker.FuncFormatter(
                    lambda x, _: _format_chart_value(x, y_format)))
                ax.grid(axis='y', alpha=0.3)
            elif is_horizontal:
                bars = ax.barh(x_values, y_values, color=bar_colors, height=0.5, edgecolor='white', linewidth=0.5)
                ax.set_xlabel(spec.get("y_label", ""), fontsize=10)
                ax.xaxis.set_major_formatter(mticker.FuncFormatter(
                    lambda x, _: _format_chart_value(x, y_format)))
                ax.grid(axis='x', alpha=0.3)
                # Value labels
                max_val = max(y_values) if y_values else 1
                for bar, val in zip(bars, y_values):
                    label = _format_chart_value(val, y_format)
                    ax.text(bar.get_width() + max_val * 0.02, bar.get_y() + bar.get_height() / 2,
                            label, ha='left', va='center', fontsize=9, fontweight='bold')
            else:
                bars = ax.bar(x_values, y_values, color=bar_colors, width=0.5, edgecolor='white', linewidth=0.5)
                ax.set_ylabel(spec.get("y_label", ""), fontsize=10)
                ax.yaxis.set_major_formatter(mticker.FuncFormatter(
                    lambda x, _: _format_chart_value(x, y_format)))
                ax.grid(axis='y', alpha=0.3)
                # Value labels
                max_val = max(y_values) if y_values else 1
                for bar, val in zip(bars, y_values):
                    label = _format_chart_value(val, y_format)
                    ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + max_val * 0.02,
                            label, ha='center', va='bottom', fontsize=9, fontweight='bold')

            ax.set_title(title, fontsize=13, fontweight='bold', pad=12)
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            ax.tick_params(labelsize=9)

            # Legend for color-keyed charts
            if color_key and isinstance(colors_config, dict):
                from matplotlib.patches import Patch
                legend_elements = [Patch(facecolor=c, label=k) for k, c in colors_config.items()]
                ax.legend(handles=legend_elements, fontsize=8, loc='upper left')

            plt.tight_layout()
            buf = BytesIO()
            fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            plt.close(fig)
            buf.seek(0)
            images.append((chart_id, buf))
        except Exception as e:
            logger.warning(f"Failed to generate chart image for '{chart_id}': {e}")
            continue

    return images


def generate_memo_docx(
    memo: Dict[str, Any],
    company_name: str = "Company"
) -> BytesIO:
    """
    Generate a DOCX file from memo data.

    Args:
        memo: Dictionary containing memo sections and metadata
        company_name: Name of the company for the cover page

    Returns:
        BytesIO buffer containing the DOCX file
    """
    logger.info(f"Generating DOCX for memo {memo.get('id', 'unknown')}")

    doc = Document()

    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set up styles
    _setup_styles(doc)

    # Add cover page
    _add_cover_page(
        doc,
        company_name=company_name,
        ticket_size=memo.get('ticket_size'),
        post_money_valuation=memo.get('post_money_valuation'),
        created_at=memo.get('created_at', '')
    )

    # Add table of contents
    _add_table_of_contents(doc)

    # Generate chart images from metadata
    chart_image_list: List[Tuple[str, BytesIO]] = []
    metadata_raw = memo.get('metadata')
    if metadata_raw:
        try:
            metadata = json.loads(metadata_raw) if isinstance(metadata_raw, str) else metadata_raw
            chart_data = metadata.get('chart_data', {})
            if chart_data:
                chart_image_list = _generate_chart_images(chart_data)
        except (json.JSONDecodeError, TypeError, AttributeError):
            pass

    # Add each section
    for i, (key, label) in enumerate(MEMO_SECTIONS, 1):
        content = memo.get(key)
        if content:
            # Section header
            doc.add_paragraph(f"{i}. {label}", style='Heading 1')

            # Section content
            _markdown_to_docx(doc, content)

            # Embed charts after Financial Analysis section
            if key == 'financial_analysis' and chart_image_list:
                doc.add_paragraph()
                for _chart_id, buf in chart_image_list:
                    doc.add_picture(buf, width=Inches(5.5))
                    doc.add_paragraph()

            # Add some spacing between sections
            doc.add_paragraph()

    # Save to BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    logger.info(f"DOCX generated successfully, size: {buffer.getbuffer().nbytes} bytes")
    return buffer


if __name__ == "__main__":
    # Test with sample memo
    sample_memo = {
        "id": "test-memo",
        "executive_summary": "## Overview\n\nThis is a **test** memo with *italic* text.\n\n- Bullet point 1\n- Bullet point 2\n\n1. Numbered item\n2. Another item",
        "market_analysis": "The market is growing rapidly.",
        "team_assessment": "Strong founding team.",
        "product_technology": "Innovative technology stack.",
        "financial_analysis": "Solid financials.",
        "risks_concerns": "Key risks identified.",
        "outcome_scenario_analysis": "Multiple scenarios analyzed.",
        "investment_recommendation": "Recommend investment.",
        "ticket_size": 500000,
        "post_money_valuation": 10000000,
        "created_at": datetime.now().isoformat()
    }

    buffer = generate_memo_docx(sample_memo, "Test Company")
    with open("test_memo.docx", "wb") as f:
        f.write(buffer.getvalue())
    print("Test DOCX saved to test_memo.docx")
