"""
Word document parsing tool (.docx and .doc) with text, table, and metadata extraction.
Uses python-docx for .docx files and falls back to textract/antiword for legacy .doc files.
"""

import sys
from pathlib import Path
from typing import Dict, List, Any
from loguru import logger

try:
    from docx import Document as DocxDocument
    from docx.opc.exceptions import PackageNotFoundError
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. Install with: pip install python-docx")


def _extract_docx(file_path: Path) -> Dict[str, Any]:
    """
    Extract text, tables, and metadata from a .docx file.

    Args:
        file_path: Path to the .docx file

    Returns:
        Dictionary with extracted content
    """
    result = {
        "file_name": file_path.name,
        "file_path": str(file_path.absolute()),
        "file_size": file_path.stat().st_size,
        "text": "",
        "tables": [],
        "metadata": {},
        "file_type": "docx",
        "method": "python-docx",
    }

    try:
        doc = DocxDocument(str(file_path))
    except PackageNotFoundError:
        result["error"] = "File is not a valid .docx document"
        logger.error(f"Invalid .docx file: {file_path}")
        return result
    except Exception as e:
        result["error"] = f"Failed to open .docx file: {str(e)}"
        logger.error(f"Failed to open .docx: {file_path}: {e}")
        return result

    # --- Extract metadata from core properties ---
    try:
        props = doc.core_properties
        result["metadata"] = {
            "title": props.title or "",
            "author": props.author or "",
            "subject": props.subject or "",
            "created": str(props.created) if props.created else "",
            "modified": str(props.modified) if props.modified else "",
            "last_modified_by": props.last_modified_by or "",
        }
    except Exception as e:
        logger.debug(f"Could not extract .docx metadata: {e}")

    # --- Extract text from paragraphs ---
    paragraphs: List[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Preserve heading structure by prefixing with markers
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading", "").strip()
                prefix = "#" * int(level) if level.isdigit() else "#"
                paragraphs.append(f"{prefix} {text}")
            else:
                paragraphs.append(text)

    # --- Extract text from tables ---
    for table_idx, table in enumerate(doc.tables):
        try:
            table_data: List[List[str]] = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)

            if table_data:
                headers = table_data[0] if table_data else []
                result["tables"].append({
                    "table_index": table_idx,
                    "rows": len(table_data),
                    "columns": len(headers),
                    "data": table_data,
                    "headers": headers,
                })

                # Also add table content as text for full-text search
                table_lines = []
                for row_data in table_data:
                    row_text = " | ".join(cell for cell in row_data if cell)
                    if row_text.strip():
                        table_lines.append(row_text)
                if table_lines:
                    paragraphs.append("\n".join(table_lines))
        except Exception as e:
            logger.warning(f"Failed to extract table {table_idx}: {e}")

    result["text"] = "\n\n".join(paragraphs)
    result["total_chars"] = len(result["text"])
    result["total_tables"] = len(result["tables"])

    return result


def _extract_doc_fallback(file_path: Path) -> Dict[str, Any]:
    """
    Extract text from a legacy .doc file using subprocess fallback (antiword or textutil on macOS).

    Args:
        file_path: Path to the .doc file

    Returns:
        Dictionary with extracted content
    """
    import subprocess
    import shutil

    result = {
        "file_name": file_path.name,
        "file_path": str(file_path.absolute()),
        "file_size": file_path.stat().st_size,
        "text": "",
        "tables": [],
        "metadata": {},
        "file_type": "doc",
        "method": "fallback",
    }

    text = ""

    # Try antiword first (cross-platform)
    if shutil.which("antiword"):
        try:
            proc = subprocess.run(
                ["antiword", str(file_path)],
                capture_output=True, text=True, timeout=60
            )
            if proc.returncode == 0 and proc.stdout.strip():
                text = proc.stdout
                result["method"] = "antiword"
        except Exception as e:
            logger.debug(f"antiword failed: {e}")

    # macOS fallback: textutil converts .doc to plain text
    if not text and shutil.which("textutil"):
        try:
            proc = subprocess.run(
                ["textutil", "-convert", "txt", "-stdout", str(file_path)],
                capture_output=True, text=True, timeout=60
            )
            if proc.returncode == 0 and proc.stdout.strip():
                text = proc.stdout
                result["method"] = "textutil"
        except Exception as e:
            logger.debug(f"textutil failed: {e}")

    # LibreOffice fallback: convert to plain text
    if not text and shutil.which("libreoffice"):
        try:
            import tempfile
            with tempfile.TemporaryDirectory() as tmp_dir:
                proc = subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", "txt:Text", "--outdir", tmp_dir, str(file_path)],
                    capture_output=True, text=True, timeout=120
                )
                if proc.returncode == 0:
                    txt_file = Path(tmp_dir) / (file_path.stem + ".txt")
                    if txt_file.exists():
                        text = txt_file.read_text(encoding="utf-8", errors="replace")
                        result["method"] = "libreoffice"
        except Exception as e:
            logger.debug(f"libreoffice conversion failed: {e}")

    if not text:
        result["error"] = (
            "Cannot parse legacy .doc file. Install antiword (apt install antiword / brew install antiword) "
            "or convert to .docx format."
        )
        logger.error(f"No .doc extraction tool available for: {file_path}")
        return result

    result["text"] = text.strip()
    result["total_chars"] = len(result["text"])
    result["total_tables"] = 0

    return result


def parse_docx(file_path: str) -> Dict[str, Any]:
    """
    Parse a Word document (.docx or .doc).

    Args:
        file_path: Path to the Word document

    Returns:
        Dictionary with parsed content matching the standard parser output format:
        {
            "file_name": str,
            "file_path": str,
            "file_size": int,
            "text": str,
            "tables": list,
            "metadata": dict,
            "file_type": "docx" | "doc",
            "method": str,
            "total_chars": int,
            "total_tables": int,
        }
    """
    file_path_obj = Path(file_path)

    if not file_path_obj.exists():
        raise FileNotFoundError(f"Word document not found: {file_path}")

    if file_path_obj.stat().st_size == 0:
        return {
            "file_name": file_path_obj.name,
            "file_path": str(file_path_obj.absolute()),
            "file_size": 0,
            "text": "",
            "tables": [],
            "metadata": {},
            "file_type": file_path_obj.suffix.lower().lstrip("."),
            "method": "none",
            "error": "File is empty (0 bytes)",
            "total_chars": 0,
            "total_tables": 0,
        }

    suffix = file_path_obj.suffix.lower()

    if suffix == ".docx":
        if not DOCX_AVAILABLE:
            return {
                "file_name": file_path_obj.name,
                "file_path": str(file_path_obj.absolute()),
                "file_size": file_path_obj.stat().st_size,
                "text": "",
                "tables": [],
                "metadata": {},
                "file_type": "docx",
                "method": "none",
                "error": "python-docx is not installed. Run: pip install python-docx",
                "total_chars": 0,
                "total_tables": 0,
            }
        logger.info(f"Parsing Word document: {file_path_obj.name}")
        result = _extract_docx(file_path_obj)
    elif suffix == ".doc":
        logger.info(f"Parsing legacy .doc document: {file_path_obj.name}")
        result = _extract_doc_fallback(file_path_obj)
    else:
        return {
            "file_name": file_path_obj.name,
            "file_path": str(file_path_obj.absolute()),
            "file_size": file_path_obj.stat().st_size,
            "text": "",
            "tables": [],
            "metadata": {},
            "file_type": suffix.lstrip("."),
            "method": "none",
            "error": f"Unsupported Word document format: {suffix}",
            "total_chars": 0,
            "total_tables": 0,
        }

    if result.get("error"):
        logger.error(f"Parse error for {file_path_obj.name}: {result['error']}")
    else:
        logger.success(
            f"Parsed {file_path_obj.name}: {result.get('total_chars', 0):,} chars, "
            f"{result.get('total_tables', 0)} tables"
        )

    return result
